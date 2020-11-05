# -*- coding: utf-8 -*-
"""
Created on Mon Jun  8 16:59:32 2020

@author: Gene764
"""
import os
import requests
from bs4 import BeautifulSoup
from urllib.parse import parse_qs, urlparse
import pandas as pd
import openpyxl
from docx import Document
from docx.shared import Inches
from time import sleep
# Google 搜尋 URL
google_url = 'https://www.google.com.tw/search'

# 查詢參數
my_params = {'q': 'movie'}

# 下載 Google 搜尋結果
r = requests.get(google_url, params = my_params)
title=[]
url=[]
img=[]
count=1
def crop_html(url_str,title):
    from selenium import webdriver     #從selenium庫匯入webdirver
    from selenium.webdriver.chrome.options import Options

    chrome_options = Options()
    chrome_options.add_argument('--headless')
    chrome_options.add_argument('--disable-gpu')
    browser = webdriver.Chrome(options=chrome_options )
    path="./image1/"
    picName=path+title+".png"
    if not os.path.exists(path):
        os.mkdir(path)
#    browser=webdriver.PhantomJS(executable_path='E:/phantomjs-2.1.1-windows/bin/phantomjs.exe')    #使用webdirver.PhantomJS()方法新建一個phantomjs的物件，這裡會使用到phantomjs.exe，環境變數path中找不到phantomjs.exe，則會報錯
#    sleep(2)
    browser.get(url_str)           #使用get()方法，開啟指定頁面。注意這裡是phantomjs是無介面的，所以不會有任何頁面顯示
 
    browser.maximize_window()      #設定phantomjs瀏覽器全屏顯示
 
    browser.save_screenshot(picName)   #使用save_screenshot將瀏覽器正文部分截圖，即使正文字分無法一頁顯示完全，save_screenshot也可以完全截圖
#    sleep(1)
    browser.close()           #關閉phantomjs瀏覽器，不要忽略了這一步，否則你會在任務瀏覽器中發現許多phantomjs程序

def getData(req):
# 確認是否下載成功
    global count
    if req.status_code == requests.codes.ok:
        
  # 以 BeautifulSoup 解析 HTML 原始碼
      soup = BeautifulSoup(req.text, 'html.parser')
  # 觀察 HTML 原始碼
  # print(soup.prettify())

  # 以 CSS 的選擇器來抓取 Google 的搜尋結果
#  items = soup.select('div.g > h3.r > a[href^="/url"]')
      items_a = soup.select('a')
#      print(items_a)
      items_h3 = soup.select('h3')
      for i in items_a:
          for j in items_h3:
              if j.text in i.text:   
                try:  
                    str_count=str(count)
                    str_count=str_count.zfill(3)
                    qs = urlparse(i.get('href')).query
                    parsed_qs = parse_qs(qs)
                    url_path=parsed_qs.get('q', [])[0]
                    title.append(j.text)
                    url.append(url_path)
                    img.append("./image1/"+str_count+".png")
                    count=count+1
                    print(str_count)
                    crop_html(url_path,str_count)
                    print("標題：" + j.text)                              
                    print("網址：" + parsed_qs.get('q', [])[0])
                    
                except:
                    pass
     
      
      page2=soup.select('a[aria-label="下一頁"]')
      r = requests.get('https://www.google.com.tw'+page2[0].get('href'))
      getData(r)

if __name__=="__main__":
    try:      
        getData(r)
        
    except:
        #寫入excel
        file_name='movie'
        savefolder='E:/project/searchGoogle/'
        writer=pd.ExcelWriter(savefolder+'{}.xlsx'.format(file_name), engine='xlsxwriter',options={'strings_to_urls': False}) 
        df1=pd.DataFrame({'title':title})
        df2=pd.DataFrame({'url':url})
        df1.to_excel(writer,sheet_name='Data1',startcol=0,index=False)
        df2.to_excel(writer,sheet_name='Data1',startcol=1,index=False)
        writer.close()
        
        #寫入word
        document = Document()  
        document.add_heading('Movie', 0)
    
    #    p = document.add_paragraph('A plain paragraph having some ')
    #    p.add_run('bold').bold = True
    #    p.add_run(' and some ')
    #    p.add_run('italic.').italic = True
    
    #    document.add_heading('Heading, level 1', level=1)
    #    document.add_paragraph('Intense quote', style='IntenseQuote')    
        for i in range(len(title)):
            try:        
                document.add_paragraph(
                    'title : '+title[i], style='ListBullet'
                )
                document.add_paragraph(
                    'url : '+url[i], style='ListBullet'
                )
                
                document.add_picture(img[i], width=Inches(5.0))
            except:
                pass
        
        document.add_page_break()
        document.save('movie.docx')
        
    #    wb = openpyxl.load_workbook(savefolder+'{}.xlsx'.format(file_name))
    #    ws = wb.worksheets[0]
    #    
    #    for i in range(134):
    #        img_name=str(i+1)
    #        img_name=img_name.zfill(3)
    #        imag = openpyxl.drawing.image.Image("./image/"+img_name+'.png')
    #        imag.width, imag.height = int(imag.width/2), int(imag.height/2)
    #    
    #        imag.anchor = 'C'+str(i+2)
    #               
    #        ws.add_image(imag)
    #    wb.save(savefolder+'{}.xlsx'.format(file_name))
    #    wb.close()